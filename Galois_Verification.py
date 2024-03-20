import sys
import keyboard
import qdarkstyle
from PyQt5.QtWidgets import (QApplication, QMainWindow, QSystemTrayIcon, QMenu, QAction, QLabel,
                            QPushButton, QTextBrowser, QCheckBox, QWidgetAction, QTextEdit, 
                            QVBoxLayout, QHBoxLayout, QWidget, QDesktopWidget, QFileDialog,
                            QShortcut)
from PyQt5.QtGui import QIcon, QPixmap, QKeySequence, QFont, QPalette, QColor
from PyQt5.QtCore import Qt, QProcess, QMetaObject, pyqtSlot, pyqtSignal
from PyQt5.QtWebEngineWidgets import QWebEngineView
from PPOCR_api import GetOcrApi
from PPOCR_visualize import visualize
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from fuzzywuzzy import fuzz, process
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import difflib
import cv2
# import tbpu
import ujson as json
import win32api
import ctypes
import re
import os
import numpy as np
import unicodedata
from datetime import datetime
# from html import escape

from pprint import pprint

# def clean_text(text):
#     # 移除特殊符号和空格
#     cleaned_text = re.sub(r'[^\w\s]', '', text)
#     # 统一转为小写
#     cleaned_text = cleaned_text.lower()
#     return cleaned_text

def calculate_cosine_similarity(text1, text2):
    vectorizer = CountVectorizer().fit_transform([text1, text2])
    cosine_similarities = cosine_similarity(vectorizer)
    return cosine_similarities[0, 1]

def find_best_match(ocr_text, reference_texts):
    if reference_texts:
        best_match = reference_texts[0]
        best_similarity = calculate_cosine_similarity(ocr_text, best_match)
    else:
        return "文档中未找到匹配内容", 0
    
    for ref_text in reference_texts[1:]:
        similarity = calculate_cosine_similarity(ocr_text, ref_text)
        if similarity > best_similarity:
            best_similarity = similarity
            best_match = ref_text
    
    return best_match, best_similarity

def find_diff_box(ocr_text, reference_texts):
    if reference_texts:
        best_match = reference_texts[0]['text']
        best_similarity = calculate_cosine_similarity(ocr_text, best_match)
        best_match_box = reference_texts[0]['box']
    else:
        return "文档中未找到匹配内容", 0, None
    
    for ref_text_data in reference_texts[1:]:
        ref_text = ref_text_data['text']
        similarity = calculate_cosine_similarity(ocr_text, ref_text)
        if similarity > best_similarity:
            best_similarity = similarity
            best_match = ref_text
            best_match_box = ref_text_data['box']
    
    return best_match, best_similarity, best_match_box

def find_best_table_index(ocr_text, table_texts):
    try:
        similarities = [calculate_cosine_similarity(ocr_text, table_text) for table_text in table_texts]
        # print(similarities)
        # print(similarities.index(max(similarities)))
        return similarities.index(max(similarities))
    except Exception as e:
        print(f"未找到相似度最高的匹配项 {e}")
    
def draw_rectangles(image_path, rectangles, text_list, color=(192, 128, 0, 255), text_background_color=(0, 128, 192, 192)):
    # 读取 Pillow Image
    image = Image.open(image_path).convert("RGBA")
    draw = ImageDraw.Draw(image)

    # 设置字体
    font_path = "SimHei.ttf"
    font_size = 16
    font = ImageFont.truetype(font_path, font_size)

    for rectangle, text in zip(rectangles, text_list):
        # 下划线策略
        start_point = (rectangle[3][0] - 2, rectangle[3][1] - 2)  # 左下角坐标
        end_point = (rectangle[2][0] - 2, rectangle[2][1] - 2)  # 右下角坐标

        # 在图像上绘制矩形框或下划线
        thickness = 1  # 线条厚度
        draw.line([start_point, end_point], fill=color, width=thickness)

        # 计算文字位置，使其位于当前划线的下方
        text_position = (start_point[0], start_point[1] + 1)  # 文字位置

        # 获取文字包围盒
        text_bbox = draw.textbbox(text_position, text, font=font)

        # 绘制带有半透明背景颜色的文字
        background_rect = [text_bbox[0], text_bbox[1], text_bbox[2], text_bbox[3]]

        # 计算透明度
        alpha = text_background_color[3] / 255.0

        # 应用透明度到文字背景颜色
        background_color_with_alpha = (
            int(text_background_color[0] * alpha),
            int(text_background_color[1] * alpha),
            int(text_background_color[2] * alpha),
            text_background_color[3]
        )

        # 获取图像的一个区域，与矩形相交
        region = image.crop(background_rect)
        # 将区域与文字背景进行融合
        blended = Image.blend(region, Image.new('RGBA', region.size, background_color_with_alpha), alpha=alpha)
        # 将融合后的图像粘贴回原图像中
        image.paste(blended, background_rect)

        # 将文字绘制到图像上
        draw.text(text_position, text, font=font, fill=(255, 160, 0, 255))

    # 保存绘制了矩形框的图像
    current_time = datetime.now()
    formatted_time = current_time.strftime("%Y%m%d%H%M%S")
    output_path = f'temp\\{formatted_time}.png'
    image.save(output_path)

    return output_path


def remove_white_background(image_path):
    # 打开图像
    image = Image.open(image_path).convert("RGBA")
    
    # 获取图像的宽度和高度
    width, height = image.size
    
    # 遍历图像的每个像素
    for x in range(width):
        for y in range(height):
            # 获取当前像素的颜色
            r, g, b, a = image.getpixel((x, y))
            
            # 如果像素是白色（r、g、b 值都是255）
            if r == 255 and g == 255 and b == 255:
                # 将白色像素设置为透明
                image.putpixel((x, y), (r, g, b, 0))
    
    # 保存处理后的图像
    output_path = 'temp\\'+'output_image.png'
    image.save(output_path, format="PNG")
    
    return output_path

def detect_lines(image, thickness=1, line_color=255):
    # 获取图像的高度和宽度
    pixel_height_line = image.shape[0]
    pixel_width_line = image.shape[1]

    # 记录横线的上下边界
    horizontal_lines = []

    # pixel_values = list(range(31))

    # 遍历图像的高度
    for height_x in range(0, pixel_height_line, thickness):
        count = 0
        # 检查横向连续的黑色像素长度
        for width_x in range(pixel_width_line - thickness + 1):
            # 如果当前位置及粗度范围内的像素都为黑色
            if np.all(np.any(image[height_x, width_x:width_x + thickness] == 0)):
                count += 1

            # 如果连续黑色像素长度超过当前粗度阈值的85%
            if count > int(pixel_width_line * 0.85):
                # 将该横向区域的像素值全部设为白色（255），去除横线
                image[height_x, :] = line_color
                # 记录横线的上下边界
                horizontal_lines.append([height_x, height_x + thickness - 1])
                break

    return image, horizontal_lines

def should_match(text):
    # 确保文本不包含英文逗号、英文句号或中文句号
    if not re.search(r'[.。]', text):
        return True
    return False

def find_similar_sentences(text1, text2, threshold=0.8):
    matches = []
    for sent1 in text1:
        for sent2 in text2:
            s = difflib.SequenceMatcher(None, sent1, sent2)
            if s.ratio() > threshold:
                matches.append((sent1, sent2))
    return matches

def extract_table_from_docx(docx_path):
    doc = Document(docx_path)
    extracted_content = []
    
    # 遍历文档中的表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extracted_content.append(cell.text.strip())
    
    return "\n".join(extracted_content)

# def normalize_text(text):
#     # 将文本标准化为Unicode规范化形式C（NFC）
#     return unicodedata.normalize('NFC', text)

def map_diff_indices(text1, text2, opcodes):
    indices_map = []
    for tag, i1, i2, j1, j2 in opcodes:
        print(f"Tag: {tag}, Indices: {i1}-{i2}, {j1}-{j2}")
        if tag == 'equal':
            for i in range(i1, i2):
                indices_map.append((i, i))
        elif tag in ('delete', 'replace'):
            for i in range(i1, i2):
                indices_map.append((i, None))
        elif tag in ('insert', 'replace'):
            for j in range(j1, j2):
                indices_map.append((None, j))
    return indices_map

def highlight_differences(text1, text2, dictionary):
    # 注意事项，取消高亮的法则只要保证`ocr_false_detection.txt`中的忽略规则冒号右边的字不碰巧为这句ocr识别出来的match[0]对应的原句match[1]中其他字就不存在这样的问题。
    # 使用difflib库进行文本比对，生成差异列表
    d = difflib.ndiff(text1, text2)
    diff_list = list(d)
    
    # 存储结果的列表，用于分别存储两段文本的高亮部分
    result1 = []
    result2 = []
    
    # 遍历差异列表
    for item in diff_list:
        # 如果是相同的部分，则直接添加到结果列表中
        if item[0] == " ":
            result1.append(item[2:])
            result2.append(item[2:])
        # 如果是text1多出来的部分，进行处理
        elif item[0] == "-":
            word = item[2:]  # 获取被删除的词语
            # 检查该词语是否在dictionary规则内
            # 如果在规则内，不高亮显示，直接添加到结果列表中
            if any(word in values[1] for values in dictionary.values()):
                result1.append(word)
            # 如果不在规则内，高亮显示
            else:
                result1.append(f'<span style="color: #fd7bd4;">{word}</span>')
        # 如果是text2多出来的部分，进行处理
        elif item[0] == "+":
            word = item[2:]  # 获取新增的词语
            # 检查该词语是否在dictionary规则内
            # 如果在规则内，不高亮显示，直接添加到结果列表中
            if any(word in values[1] for values in dictionary.values()):
                result2.append(word)
            # 如果不在规则内，高亮显示
            else:
                result2.append(f'<span style="color: #55a6f7;">{word}</span>')
    
    # 将结果列表转换为字符串，用于在HTML页面中显示
    return "".join(result1), "".join(result2)


# def highlight_differences(text1, text2):
#     # 排错使用版本
#     print("Input Text1:", text1)
#     print("Input Text2:", text2)
#     d = difflib.ndiff(text1, text2)
#     diff_list = list(d)
    
#     result1 = []
#     result2 = []
#     for item in diff_list:
#         if item[0] == " ":
#             result1.append(item[2:])
#             result2.append(item[2:])
#         elif item[0] == "-":
#             result1.append(f'<span style="background-color: #fd7bd4;">{item[2:]}</span>')
#         elif item[0] == "+":
#             result2.append(f'<span style="background-color: #55a6f7;">{item[2:]}</span>')
    
#     print("Result1:", "".join(result1))
#     print("Result2:", "".join(result2))
    
#     return "".join(result1), "".join(result2)


scrollbar_style = """
<style>
    ::-webkit-scrollbar {
        width: 10px;  /* 设置滚动条宽度 */
    }
    ::-webkit-scrollbar-track {
        background: #19232d;  /* 设置滚动条轨道的颜色 */
        /* border: 1px solid #455364; */  /* 设置滚动条轨道的描边 */
        border-radius: 5px;   /* 设置滚动条轨道的圆角 */
        /* box-shadow: inset 0 0 5px grey; */
    }
    ::-webkit-scrollbar-thumb {
        background: #60798b;  /* 设置滚动条的颜色 */
        border-radius: 5px;   /* 设置滚动条的圆角 */
    }
    ::-webkit-scrollbar-thumb:hover {
        background: #346792;  /* 设置鼠标悬停时滚动条的颜色 */
    }
</style>
"""

class MainApp(QMainWindow):

    triggerScreenshot = pyqtSignal()  # 定义一个信号

    def __init__(self):
        super(MainApp, self).__init__()

        self.init_ui()
        self.init_tray_icon()
        self.init_hotkeys()

        self.triggerScreenshot.connect(self.take_screenshot)  # 连接信号到take_screenshot函数

        self.false_detection_rules = self.load_false_detection_rules()
        # print(self.false_detection_rules)

        # 初始化历史堆栈和当前索引
        self.history = [""]
        self.current_index = 0

        # 表单模式
        self.is_table_mode = 0
    
    def load_false_detection_rules(self):
        rules = {}
        try:
            with open('ocr_false_detection.txt', 'r', encoding='utf-8') as file:
                for line in file:
                    # 分割差异字符和忽略字符
                    diff_char, ignore_chars = line.strip().split(':')
                    # 如果忽略字符不为空，分割成列表
                    ignore_chars = ignore_chars.split(',')
                    # 替换"null"为""
                    ignore_chars = [char if char != 'null' else '' for char in ignore_chars]
                    # 将规则添加到字典中
                    rules[diff_char] = ignore_chars
        except FileNotFoundError:
            print("OCR误检测规则文件不存在")
            return {}
        except ValueError:
            # 如果文件存在但是格式错误，打印错误信息
            print("OCR误检测规则文件格式错误")
            return {}
        return rules

    def init_ui(self):
        """初始化主要的用户界面组件"""
        self.setWindowTitle("Galois's Content Verification")
        self.setGeometry(100, 100, 1000, 800)
        # self.setFixedSize(1000, 800)
        self.center()

        # 创建小部件
        self.label_top = QLabel("———————————————————— Welcome to Galois's Content Verification machine for Product ————————————————————", self)
        font = QFont()
        font.setPointSize(11)
        font.setBold(True)
        font.setFamily("Arial")
        self.label_top.setFont(font)
        palette = self.label_top.palette()
        palette.setColor(QPalette.WindowText, QColor("#FF5733"))
        self.label_top.setPalette(palette)
        self.label_top.setAlignment(Qt.AlignCenter)

        self.results_browser = QWebEngineView(self)
        
        display_browser_font = QFont()
        display_browser_font.setPointSize(8)
        display_browser_font.setBold(True)
        display_browser_font.setFamily("Arial")

        vbox_display = QVBoxLayout()
        label_display = QLabel("OCR Display Area", self)
        label_display.setFont(display_browser_font)
        label_display.setStyleSheet("background-color: #101c28; border: 1px solid #555555; border-bottom: none; padding: 2px; border-top-left-radius: 10px; border-top-right-radius: 10px;")
        self.text_display = QTextBrowser(self)
        self.text_display.setStyleSheet("border-top: none; border-top-left-radius: 0px; border-top-right-radius: 0px;")
        vbox_display.addWidget(label_display)
        vbox_display.addWidget(self.text_display)
        self.button_comparison = QPushButton("文本校验", self)
        self.button_form_comparison = QPushButton("表单校验", self)

        hbox_buttons = QHBoxLayout()
        hbox_buttons.addWidget(self.button_comparison)
        hbox_buttons.addWidget(self.button_form_comparison)

        hbox_buttons.setSpacing(5)
        hbox_buttons.setContentsMargins(0, 5, 0, 5)

        vbox_display.addLayout(hbox_buttons)

        vbox_display.setSpacing(0)
        vbox_display.setContentsMargins(0, 0, 0, 0)
        
        button_select = QPushButton("选择参考文档", self)
        # 新建切换按钮并初始化为禁用表格模式
        self.toggle_table_mode_button = QPushButton("文本流模式", self)
        self.toggle_table_mode_button.setCheckable(True)

        hbox_select = QHBoxLayout()
        hbox_select.addWidget(self.toggle_table_mode_button)

        hbox_select.setSpacing(5)
        hbox_select.setContentsMargins(0, 5, 0, 5)

        hbox_select.addWidget(button_select)

        self.text_browser = QTextBrowser(self)
        self.text_browser.setStyleSheet("border-top: none; border-top-left-radius: 0px; border-top-right-radius: 0px;")
        # self.text_browser = QWebEngineView(self)
        
        vbox_browser = QVBoxLayout()
        label_browser = QLabel("Original Content", self)
        label_browser.setFont(display_browser_font)
        label_browser.setStyleSheet("background-color: #101c28; border: 1px solid #555555; border-bottom: none; padding: 2px; border-top-left-radius: 10px; border-top-right-radius: 10px;")
        vbox_browser.addWidget(label_browser)
        vbox_browser.addWidget(self.text_browser)
        vbox_browser.addLayout(hbox_select)

        vbox_browser.setSpacing(0)
        vbox_browser.setContentsMargins(0, 0, 0, 0)

        self.results_browser.page().setBackgroundColor(Qt.transparent)
        self.results_browser.setStyleSheet("border: none;")
        self.results_browser.setMinimumHeight(100)
        self.results_browser.setStyleSheet("background-color: #19232d;")

        # 布局
        hbox_inner = QHBoxLayout()
        hbox_inner.addLayout(vbox_display)
        hbox_inner.addLayout(vbox_browser)

        vbox_outer = QVBoxLayout()
        vbox_outer.addWidget(self.label_top)
        vbox_outer.addLayout(hbox_inner)
        vbox_outer.addWidget(self.results_browser, 1)

        central_widget = QWidget(self)
        central_widget.setLayout(vbox_outer)
        self.setCentralWidget(central_widget)

        self.setWindowIcon(QIcon("icon.ico"))
        self.label_top.setFocus()

        button_select.clicked.connect(self.select_reference_file)
        self.button_comparison.clicked.connect(self.compare_texts)
        # self.button_form_comparison.clicked.connect(self.on_form_comparison_clicked)
        self.button_form_comparison.clicked.connect(self.on_compare_button_clicked)
        self.toggle_table_mode_button.toggled.connect(self.toggle_table_mode)

        # 默认状态
        self.button_comparison.setEnabled(False)
        self.button_form_comparison.setEnabled(False)
        self.toggle_table_mode_button.setEnabled(False)
        button_select.setEnabled(True)

    def init_tray_icon(self):
        """初始化系统托盘图标及其关联的操作和菜单。"""
        self.tray_icon = QSystemTrayIcon(self)
        self.tray_icon.setIcon(QIcon("icon.ico"))

        self.tray_menu = QMenu()
        self.show_action = QAction("显示主界面", self)
        self.hide_when_screenshot = QAction("截图时隐藏窗口", self)
        self.diff_show = QAction("仅展示差异内容", self)
        self.hide_when_screenshot.setCheckable(True)
        self.diff_show.setCheckable(True)
        self.diff_show.setChecked(True)
        self.hide_tolerance_action = QAction("不展示OCR误差", self)
        self.hide_tolerance_action.setCheckable(True)
        self.hide_tolerance_action.setChecked(True)  # 初始状态为选中
        self.show_image = QAction('不展示图像分析', self)
        self.show_image.setCheckable(True)
        self.show_image.setChecked(True)  # 初始状态为选中
        self.quit_action = QAction("退出", self)

        self.show_action.triggered.connect(self.show_main_window)
        self.quit_action.triggered.connect(self.quit_app)

        self.tray_menu.addAction(self.show_action)
        self.tray_menu.addAction(self.hide_when_screenshot)
        self.tray_menu.addAction(self.diff_show)
        self.tray_menu.addAction(self.hide_tolerance_action)
        self.tray_menu.addAction(self.show_image)
        self.tray_menu.addSeparator()
        self.tray_menu.addAction(self.quit_action)

        self.tray_icon.setContextMenu(self.tray_menu)
        self.tray_icon.show()
        self.tray_icon.activated.connect(self.tray_icon_activated)

    def init_hotkeys(self):
        """使用keyboard模块初始化热键。"""
        keyboard.add_hotkey('ctrl+d', lambda: self.triggerScreenshot.emit())
        # keyboard.add_hotkey('win+shift+a', lambda: QMetaObject.invokeMethod(self, "take_screenshot", Qt.QueuedConnection))
        # keyboard.add_hotkey('ctrl+r', self.clear_text_display)
        # keyboard.add_hotkey('ctrl+r', lambda: QMetaObject.invokeMethod(self, "clear_text_display", Qt.QueuedConnection))
        self.shortcut_clear = QShortcut(QKeySequence("Ctrl+R"), self)
        self.shortcut_clear.activated.connect(self.clear_text_display)

        # 设置快捷键
        self.shortcut_undo = QShortcut(QKeySequence("Ctrl+Z"), self)
        self.shortcut_undo.activated.connect(self.undo)
        
        self.shortcut_redo = QShortcut(QKeySequence("Ctrl+Shift+Z"), self)
        self.shortcut_redo.activated.connect(self.redo)

        self.shortcut_clear_results = QShortcut(QKeySequence("Ctrl+L"), self)  # Ctrl+L 用于清除结果
        self.shortcut_clear_results.activated.connect(self.clear_results_browser)
    
    def toggle_table_mode(self, checked):
        if checked:
            self.toggle_table_mode_button.setText("表单模式")
            self.toggle_table_mode_button.setStyleSheet("background-color: #0066aa;")
            self.is_table_mode = 1
            self.button_comparison.setEnabled(False)
            self.button_form_comparison.setEnabled(True)
        else:
            self.toggle_table_mode_button.setText("文本流模式")
            self.toggle_table_mode_button.setStyleSheet("")
            self.is_table_mode = 0
            self.button_comparison.setEnabled(True)
            self.button_form_comparison.setEnabled(False)


    def clear_results_browser(self):
        # self.results_browser.clear()
        self.results_browser.setHtml("")

    def tray_icon_activated(self, reason):
        # 托盘图标双击时，显示主界面
        if reason == QSystemTrayIcon.DoubleClick:
            self.show_main_window()
    
    def center(self):
        # 获取屏幕的矩形和窗口的矩形
        screen = QDesktopWidget().screenGeometry()
        window = self.geometry()
        self.move(int((screen.width() - window.width()) / 2), int((screen.height() - window.height()) / 2))

    def show_main_window(self):
        # 激活并显示主界面
        self.setWindowState(self.windowState() & ~Qt.WindowMinimized | Qt.WindowActive)
        self.show()
        self.activateWindow()
        self.raise_()
    
    def show_image_viewer(self, image_path):
        # 创建一个新窗口来显示图像
        self.image_viewer = QWidget()
        self.image_viewer.setWindowTitle('image view')
        self.image_viewer.setWindowIcon(QIcon("icon.ico"))
        self.image_viewer.setGeometry(100, 100, 800, 600)

        # 创建一个用于显示图像的QLabel
        image_label = QLabel()
        image_label.setAlignment(Qt.AlignCenter)  # 水平和垂直居中显示图像

        # 使用QPixmap加载图像
        pixmap = QPixmap(image_path)
        image_label.setPixmap(pixmap)

        # 创建一个垂直布局，并将图像标签添加到其中
        layout = QVBoxLayout()
        layout.addWidget(image_label)
        self.image_viewer.setLayout(layout)

        self.image_viewer.show()

    def quit_app(self):
        """退出应用程序"""
        # 关闭主窗口
        self.close()

        # 如果截图进程存在，结束该进程
        if hasattr(self, 'screenshot_process'):
            self.screenshot_process.terminate()
            self.screenshot_process.waitForFinished(2000)  # 等待2秒以给进程结束的时间

        # 隐藏托盘图标并退出应用
        self.tray_icon.hide()
        QApplication.quit()

    def closeEvent(self, event):
        # 重写窗口的关闭事件，使其隐藏窗口而不是退出应用
        event.ignore()
        self.hide()

    def update_history(self, text):
        # print("Before undo:", self.current_index, self.history)
        # 当文本更改时，将其添加到历史堆栈
        self.current_index += 1
        self.history = self.history[:self.current_index]
        self.history.append(text)
        # print("After undo:", self.current_index, self.history)
    
    def undo(self):
        # print("Before undo:", self.current_index, self.history)
        if self.current_index > 0:
            self.current_index -= 1
            self.text_display.setPlainText(self.history[self.current_index])
        # print("After undo:", self.current_index, self.history)

    def redo(self):
        # print("Before undo:", self.current_index, self.history)
        if self.current_index < len(self.history) - 1:
            self.current_index += 1
            self.text_display.setPlainText(self.history[self.current_index])
        # print("After undo:", self.current_index, self.history)
    
    def compare_texts(self):
        try:
            text_display_content = self.text_display.toPlainText()
            text_display_list = text_display_content.split('\n')
            
            # # from pprint import pprint
            # pprint(text_display_list)
            # # pprint(text_paras_list)
            # exit()
            self.para_result_matched = []
            for ocr_text in self.ocr_para:
                best_match, similarity = find_best_match(ocr_text, self.para)
                if similarity != 0:
                    self.para_result_matched.append((ocr_text, best_match, similarity))
            
            self.match_para = [item for item in self.para_result_matched if item[2] >= 0.3]
            
            # matches = find_similar_sentences(text_display_list, text_paras_list)

            result_texts = []
            for match in self.match_para:
                # 获取匹配的句子
                matched_sentence = match[1]

                # 遍历句子中的每个字符，检查是否在规则字典中
                dictionary = {}
                for start_index in range(len(matched_sentence)):
                    for end_index in range(start_index + 1, len(matched_sentence) + 1):
                        substring = matched_sentence[start_index:end_index]
                        if substring in self.false_detection_rules:
                            if len(substring) == 1:
                                if substring not in dictionary:
                                    dictionary[substring] = [[start_index], self.false_detection_rules[substring]]
                                else:
                                    dictionary[substring][0].append(start_index)
                            else:
                                if substring not in dictionary:
                                    dictionary[substring] = [[(start_index, end_index - 1)], self.false_detection_rules[substring]]
                                else:
                                    dictionary[substring][0].append((start_index, end_index - 1))

                diff_text1, diff_text2 = highlight_differences(match[0].replace(" ", ""), match[1], dictionary)

                if self.diff_show.isChecked():
                    s = difflib.SequenceMatcher(None, diff_text1, diff_text2)
                    if s.ratio() < 1.0:
                        if self.hide_tolerance_action.isChecked():
                            pattern = r'<span style="color: #55a6f7;">(.*?)</span>'
                            highlighted_texts = re.findall(pattern, diff_text2)
                            if '<span style="color: #fd7bd4;">' not in diff_text1 and all(text in dictionary for text in highlighted_texts):
                                pass
                            else:
                                result_text = (f'<div style="color: white; font-size: 13px;">'
                                                f'<hr style="border: 0; height: 1px; background-color: #22aacc;">'
                                                f'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px; margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span> {diff_text1}<br>'
                                                f'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px; margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span> {diff_text2}<br>'
                                                f'</div>')
                                result_texts.append(result_text)
                        else:
                            result_text = (f'<div style="color: white; font-size: 13px;">'
                                                f'<hr style="border: 0; height: 1px; background-color: #22aacc;">'
                                                f'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px; margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span> {diff_text1}<br>'
                                                f'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px; margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span> {diff_text2}<br>'
                                                f'</div>')
                            result_texts.append(result_text)
                else:
                    result_text = (f'<div style="color: white; font-size: 13px;">'
                                    f'<hr style="border: 0; height: 1px; background-color: #22aacc;">'
                                    f'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px; margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span> {diff_text1}<br>'
                                    f'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px; margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span> {diff_text2}<br>'
                                    f'</div>')
                    result_texts.append(result_text)

            final_html = '<html><body>' + scrollbar_style + ''.join(result_texts) + '</body></html>'
            self.results_browser.setHtml(final_html)

        except ValueError as ve:
            # 捕获特定的 ValueError 异常
            print(f"ValueError occurred: {ve}")

        except Exception as e:
            # 捕获异常并进行处理，例如打印错误信息或者记录日志
            print(f"An error occurred: {e}")
    
    def on_form_comparison_clicked(self):
        text_display_content = self.text_display.toPlainText()
        # 判断内容中是否包含换行符
        if "\n" in text_display_content:
            text_display_list = text_display_content.split("\n")
            matches = find_similar_sentences(text_display_list, self.table_list)
        else:
            matches = find_similar_sentences([text_display_content], self.table_list)

        result_texts = []
        for match in matches:
            diff_text1, diff_text2 = highlight_differences(match[0], match[1])
            if self.diff_show.isChecked():
                # 检查是否真的存在差异
                s = difflib.SequenceMatcher(None, diff_text1, diff_text2)
                if s.ratio() < 1.0:  # 如果文本不完全相同（即存在差异）
                    result_text = (f'<div style="color: white; font-size: 13px;">'
                                f'<hr style="border: 0; height: 2px; background-color: #22aacc;">'
                                f'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px; margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span> {diff_text1}<br>'
                                f'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px; margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span> {diff_text2}<br>'
                                f'</div>')
                    result_texts.append(result_text)
            else:
                result_text = (f'<div style="color: white; font-size: 13px;">'
                                f'<hr style="border: 0; height: 2px; background-color: #22aacc;">'
                                f'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px; margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span> {diff_text1}<br>'
                                f'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px; margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span> {diff_text2}<br>'
                                f'</div>')
                result_texts.append(result_text)

        final_html = '<html><body>' + scrollbar_style + ''.join(result_texts) + '</body></html>'
        self.results_browser.setHtml(final_html)
    
    def compare_and_display_results(self):
        result_texts = []
        for match in self.table_result_matched:
            # 获取匹配的句子
            matched_sentence = match[1]

            # 遍历句子中的每个字符，检查是否在规则字典中
            dictionary = {}
            for start_index in range(len(matched_sentence)):
                for end_index in range(start_index + 1, len(matched_sentence) + 1):
                    substring = matched_sentence[start_index:end_index]
                    if substring in self.false_detection_rules:
                        if len(substring) == 1:
                            if substring not in dictionary:
                                dictionary[substring] = [[start_index], self.false_detection_rules[substring]]
                            else:
                                dictionary[substring][0].append(start_index)
                        else:
                            if substring not in dictionary:
                                dictionary[substring] = [[(start_index, end_index - 1)], self.false_detection_rules[substring]]
                            else:
                                dictionary[substring][0].append((start_index, end_index - 1))
            # if dictionary != {}:
            #     print(dictionary)
            # else:
            #     print("这句话中不存在可能会被误识别的字符！")

            # d = difflib.SequenceMatcher(None, match[0].replace(" ", ""), match[1])
            # opcodes = d.get_opcodes()
            # indices_map = map_diff_indices(match[0], match[1], opcodes)
            diff_text1, diff_text2 = highlight_differences(match[0].replace(" ", ""), match[1], dictionary)
            # print(diff_text1)
            # print(diff_text2)
            # diff_text1, diff_text2 = highlight_differences(match[0], match[1], dictionary)
            if self.diff_show.isChecked():
                s = difflib.SequenceMatcher(None, diff_text1, diff_text2)
                if s.ratio() < 1.0:
                    if self.hide_tolerance_action.isChecked():
                        # if '<span style="color: #fd7bd4;">' not in diff_text1 and any('<span style="color: #55a6f7;">{}</span>'.format(key) in diff_text2 and key not in dictionary for key in dictionary.keys()):
                        # # if '<span style="color: #fd7bd4;">' not in diff_text1 and all('<span style="color: #55a6f7;">{}</span>'.format(key) not in diff_text2 for key in dictionary.keys()):
                        #     # diff_text1中不存在高亮且diff_text2中存在dictionary字典中的键名以外的文本被高亮
                        #     result_text = (f'<div style="color: white; font-size: 13px;">'
                        #                     f'<hr style="border: 0; height: 1px; background-color: #22aacc;">'
                        #                     f'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px; margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span> {diff_text1}<br>'
                        #                     f'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px; margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span> {diff_text2}<br>'
                        #                     f'</div>')
                        #     result_texts.append(result_text)
                        # elif '<span style="color: #fd7bd4;">' not in diff_text1 and all('<span style="color: #55a6f7;">{}</span>'.format(key) in diff_text2 for key in dictionary.keys()):
                        #     # diff_text1中不存在高亮且diff_text2中所有被高亮的文本都属于dictionary的键名
                        #     pass
                        # elif '<span style="color: #fd7bd4;">' not in diff_text1 and '<span style="color: #55a6f7;">' not in diff_text2:
                        #     pass
                        # else:
                        #     pass
                        pattern = r'<span style="color: #55a6f7;">(.*?)</span>'
                        highlighted_texts = re.findall(pattern, diff_text2)
                        if '<span style="color: #fd7bd4;">' not in diff_text1 and all(text in dictionary for text in highlighted_texts):
                            pass
                        else:
                            result_text = (f'<div style="color: white; font-size: 13px;">'
                                            f'<hr style="border: 0; height: 1px; background-color: #22aacc;">'
                                            f'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px; margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span> {diff_text1}<br>'
                                            f'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px; margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span> {diff_text2}<br>'
                                            f'</div>')
                            result_texts.append(result_text)
                    else:
                        result_text = (f'<div style="color: white; font-size: 13px;">'
                                            f'<hr style="border: 0; height: 1px; background-color: #22aacc;">'
                                            f'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px; margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span> {diff_text1}<br>'
                                            f'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px; margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span> {diff_text2}<br>'
                                            f'</div>')
                        result_texts.append(result_text)
            else:
                result_text = (f'<div style="color: white; font-size: 13px;">'
                                f'<hr style="border: 0; height: 1px; background-color: #22aacc;">'
                                f'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px; margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span> {diff_text1}<br>'
                                f'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px; margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span> {diff_text2}<br>'
                                f'</div>')
                result_texts.append(result_text)

# if self.hide_tolerance_action.isChecked():
#     如果diff_text1中不存在高亮且diff_text2中存在dictionary字典中的键名以外的文本被高亮，则
#         result_text = ...
#     又如果diff_text1中不存在高亮且diff_text2中不存在dictionary字典中的键名以外的文本被高亮，则pass
# else: result_text = ...

        final_html = '<html><body>' + scrollbar_style + ''.join(result_texts) + '</body></html>'
        self.results_browser.setHtml(final_html)

        pattern = r'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px; margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span>(.*?)<br>'
        diff_text1_list = re.findall(pattern, final_html, re.DOTALL)
        # for text in diff_text1_list:
        #     print(text)
        
        diff_text1_clear = [re.sub(r'<.*?>', '', text).replace(" ", "") for text in diff_text1_list]
        # pprint(diff_text1_clear)

        pattern = r'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px; margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span>(.*?)<br>'
        diff_text2_list = re.findall(pattern, final_html, re.DOTALL)

        diff_text2_clear = [re.sub(r'<.*?>', '', text) for text in diff_text2_list]

        rectangles = []
        text_list = []
        for diff1_text, diff2_text in zip(diff_text1_clear, diff_text2_clear):
            best_match, similarity, best_match_box = find_diff_box(diff1_text, self.data)
            rectangles.append(best_match_box)
            text_list.append(diff2_text)

        image_rectangles = draw_rectangles(self.image_path_after_write, rectangles, text_list)

        if self.show_image.isChecked():
            pass
        else:
            self.show_image_viewer(image_rectangles)
            # self.show_image_viewer(transparent)


    def on_compare_button_clicked(self):
        # text_display_content = self.text_display.toPlainText()
        # 假设 self.table_result_matched 已经在其他地方被填充了
        try:
            self.compare_and_display_results()
        except AttributeError:
            print("对象没有table_result_matched属性")
        except Exception as e:
            print(e)
        
    @pyqtSlot()
    def select_reference_file(self):
        try:
            options = QFileDialog.Options()
            fileName, _ = QFileDialog.getOpenFileName(self, "选择参考文档", "", "word文件 (*.docx);;所有文件 (*);;文本文件 (*.txt);;PDF 文件 (*.pdf)", options=options)
            
            if fileName:

                file_name = os.path.basename(fileName)

                # 将文件名设置为 QLabel 的文本
                self.label_top.setText(f"——————————————{file_name}——————————————")

                content_to_display = []
                self.para_list = []  # 保存符合正则表达式的段落内容
                # self.table_list = []  # 保存表格的单元格内容
                self.para_stream_list = [] # 保存文本流内容

                # 如果是 .docx 文件
                if fileName.endswith(".docx"):
                    doc = Document(fileName)

                    table_count = len(doc.tables)
                    self.table_list = [[] for _ in range(table_count)]

                    # pattern = r'^【.+?】(?:.+)'
                    pattern = r'^【.+?】(?:[^\n]+)?'

                    para_exist = False
                    table_exist = False

                    # 获取文档中的所有表格对象
                    tables = doc.tables

                    for element in doc._element.body:
                        # 如果是段落
                        if element.tag.endswith('p'):
                            para = [p for p in doc.paragraphs if p._element is element][0]
                            if re.match(pattern, para.text):
                                self.para_list.append(para.text)
                            else:
                                self.para_stream_list.append(para.text)
                                # print(para.text)
                            if para.text.strip() != '':
                                content_to_display.append(para.text)
                                para_exist = True

                        # 如果是表格
                        elif element.tag.endswith('tbl'):
                            # table = [t for t in doc.tables if t._element is element][0]
                            table_index = tables.index([t for t in tables if t._element is element][0])
                            table = tables[table_index]
                            for row in table.rows:
                                for cell in row.cells:
                                    if cell.text.strip() != '':
                                        content_to_display.append(cell.text)
                                        # self.table_list.append(cell.text)  # 保存表格单元格内容
                                        self.table_list[table_index].append(cell.text)
                                        table_exist = True
                        
                    self.para_stream_list = [sentence.strip() for sentence in self.para_stream_list if sentence.strip()]

                    self.para = []
                    for para in self.para_stream_list:
                        sentences = re.split(r'(?<=[。！？；;])', para)
                        sentences = [sentence.strip() for sentence in sentences if sentence.strip()]
                        self.para.extend(sentences)
                    self.para += self.para_list
                    
                    # pprint(self.para)

                    if not para_exist:
                        print(f"{fileName}中不存在段落内容")
                    if not table_exist:
                        print(f"{fileName}中不存在表格")

                    # print(self.table_list)
                    # pprint(self.para_stream_list)

                # 如果是其他类型的文件，例如 .txt
                else:
                    with open(fileName, 'r', encoding="utf-8") as file:
                        content_to_display.append(file.read())
                
                # from pprint import pprint
                # pprint(self.para_stream_list)
                # for para in self.para_stream_list:
                #     if len(para) >= 5:
                #         start_chars = para[:4]  # 前五个字
                #         end_chars = para[-5:]   # 后五个字
                #         print(f"-: {start_chars}, -: {end_chars}")
                #     else:
                #         pass

                # 更新 self.text_browser
                # pprint('\n'.join(content_to_display))
                self.text_browser.setText('\n'.join(content_to_display))

                if self.text_browser.toPlainText():
                    self.toggle_table_mode_button.setEnabled(True)
                    self.button_comparison.setEnabled(True)

        except Exception as e:
            print("无效的文件类型:", str(e))

    def clear_text_display(self):
        # 将当前状态保存到历史堆栈中
        self.update_history(self.text_display.toPlainText())
        self.text_display.clear()

    def take_screenshot(self):
        # 隐藏窗口，然后启动截图进程
        if self.hide_when_screenshot.isChecked():
            self.hide()
        self.start_screenshot_process()

    def start_screenshot_process(self):
        """启动截图进程的函数。"""

        # 创建一个QProcess对象，用于执行外部进程
        self.screenshot_process = QProcess(self)
        # 当有数据写入到标准输出时，连接到on_screenshot_output函数
        self.screenshot_process.readyReadStandardOutput.connect(self.on_screenshot_output)
        # 当截图进程完成时，连接到on_screenshot_done函数进行处理
        self.screenshot_process.finished.connect(self.on_screenshot_done)
        # 确定应用路径
        app_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))

        if getattr(sys, 'frozen', False):
            # 如果应用是打包的
            # screenshot_tool_path = os.path.join(app_path, 'screenshot_tool')
            screenshot_tool_path = os.path.join(app_path, 'screenshot_tool')
            screenshot_tool_path = f'"{screenshot_tool_path}"'  # 将路径用引号括起来
            self.screenshot_process.start(screenshot_tool_path)
        else:
            # 在开发环境中
            self.screenshot_process.start('python', [os.path.join(app_path, 'screenshot_tool.py')])

    def on_screenshot_output(self):
        """当screenshot_tool.py有标准输出时的处理函数"""

        # 从进程的标准输出中读取文件名
        filename = bytes(self.screenshot_process.readAllStandardOutput()).decode().strip()
        # print("Screenshot saved to:", filename)
        filename = f"{os.path.dirname(os.path.abspath(__file__))}\\{filename}"
        
        # 根据截图内容类型调用ppocr_to_data
        if self.is_table_mode:  # 判断当前是在处理表格模式还是文本模式
            self.image_path_after_write, self.data= self.ppocr_to_data(filename=filename, is_table_cell=True)
        else:
            self.image_path_after_write, self.data= self.ppocr_to_data(filename=filename, is_table_cell=False)
    
    def ppocr_to_data(self, filename, is_table_cell=False):
        try:
            ImagePath = filename
            # 初始化识别器
            # ocr = GetOcrApi(os.path.abspath(os.path.join("ppocr", "PaddleOCR-json.exe"))) # 开发环境使用

            app_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
            ocr_exe_path = os.path.join(app_path, "ppocr", "PaddleOCR-json.exe")
            ocr = GetOcrApi(ocr_exe_path) # 打包环境使用

            # print(f'初始化OCR成功，进程号为{ocr.ret.pid}')
            # print(f'\n图片路径：{ImagePath}')

            if is_table_cell:
                # import matplotlib.pyplot as plt
                image = cv2.imread(ImagePath)
                gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

                # fig, axes = plt.subplots(nrows=1, ncols=3, figsize=(15, 5))
                # # 显示原始图像
                # axes[0].imshow(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
                # axes[0].set_title("Original Image")

                gray_image[gray_image < 80] = 0
                gray_image[gray_image > 210] = 255

                # axes[1].imshow(gray_image, cmap='gray')
                # axes[1].set_title("Processed Image")

                _, gray_image = cv2.threshold(gray_image, 180, 255, cv2.THRESH_BINARY)
                processed_image, horizontal_lines = detect_lines(gray_image.copy(), thickness=1, line_color=255)
                processed_image_vertical, vertical_lines = detect_lines(gray_image.copy().T, thickness=1, line_color=255)
                final_image = cv2.bitwise_or(processed_image, processed_image_vertical.T)
                cv2.imwrite(ImagePath, final_image)

                # axes[2].imshow(final_image, cmap='gray')
                # axes[2].set_title("Final Image")

                # plt.tight_layout()

                # # 显示图形
                # plt.show()

                res = ocr.run(ImagePath)
                data = res['data']
            else:
                res = ocr.run(ImagePath)
                data = res['data']

            ocr_list = [item['text'] for item in data]
            
            if is_table_cell:  # 如果是表格单元格的内容
                # processed_text = self.process_ocr_text(ocr_list)
                # 与self.table_list中的内容进行比对
                # matches = self.compare_with_reference([processed_text])
                # TODO: 根据matches进行进一步的处理，比如高亮等
                # self.text_display.append(processed_text)  # 假设显示表格内容的控件叫做text_display
                # from pprint import pprint
                # pprint(ocr_list)
                # pprint(self.table_list)
                # pprint(matches)

                # combined_list = []
                # is_single_color = False
                # temp_text = ""
                # for text in ocr_list:
                #     # 使用正则表达式检查是否以"单"或"双"开头
                #     if re.match(r'^单', text):
                #         is_single_color = True
                #         temp_text = text
                #     elif re.match(r'^双', text) and is_single_color:
                #         combined_text = temp_text + text
                #         combined_list.append(combined_text)
                #         is_single_color = False
                #     else:
                #         combined_list.append(text)

                # 在获取ocr_list之后，对每个文本进行清洗和格式化
                # ocr_list_cleaned = [clean_text(text) for text in ocr_list]

                # matched_results = self.compare_with_reference(ocr_list_cleaned)
                self.table_result_matched = []

                table_index = find_best_table_index(ocr_list[0], [text for table in self.table_list for text in table])
                # tables_len = [0] + [len(_) for _ in self.table_list]
                tables_len = [0]
                for table in self.table_list:
                    end_index = tables_len[-1] + len(table)
                    tables_len.append(end_index)

                # print(table_index)
                # print(tables_len)

                # 找到 table_index 在 tables_len 中的哪个区间
                table_list_index = -1
                for i in range(len(tables_len) - 1):
                    if tables_len[i] <= table_index < tables_len[i + 1]:
                        table_list_index = i
                        break

                if table_list_index != -1:
                    for ocr_text in ocr_list:
                        best_match, similarity = find_best_match(ocr_text, self.table_list[table_list_index])
                        if similarity != 0:
                            self.table_result_matched.append((ocr_text, best_match, similarity))
                        # print(self.table_result_matched)
                # pprint(self.table_result_matched)
                text_content = "\n".join(item for item in ocr_list)
                self.text_display.append(text_content)
                self.update_history(self.text_display.toPlainText())

            else:
                # text_browser_content = self.text_browser.toPlainText()
                # text_browser_list = text_browser_content.split('\n')

                # matched_results = self.compare_with_reference(ocr_list)
                # matched_texts = [item[0] for item in matched_results]
                # filtered_list = [text for text in ocr_list if text not in matched_texts]
                # merged_list = []
                # current_chunk = ""
                # # pattern = r'^【.+?】(?:.+)'
                # pattern = r'^【.+?】(?:[^\n]+)?'
                # for text in filtered_list:
                #     if current_chunk and current_chunk[-1] == '。':
                #         # 如果当前块以句号结尾，则将其添加到merged_list中并清空current_chunk
                #         merged_list.append(current_chunk)
                #         current_chunk = text
                #     elif re.match(pattern, current_chunk):
                #         merged_list.append(current_chunk)
                #         current_chunk = text
                #     else:
                #         # 否则，将text追加到current_chunk中
                #         current_chunk += text

                # # 添加最后的块（如果存在）
                # if current_chunk:
                #     merged_list.append(current_chunk)

                # text_content = "\n".join(item for item in matched_texts + merged_list)
                # self.text_display.append(text_content)
                # # 将当前状态保存到历史堆栈中
                # self.update_history(self.text_display.toPlainText())
                
                self.ocr_para = []
                pattern_para = []

                pattern = r'^【.+?】(?:[^\n]+)?'
                for i in range(len(ocr_list)):
                    match = re.match(pattern, ocr_list[i])
                    if match:
                        # 将符合正则表达式的部分添加到 pattern_para 中
                        pattern_para.append(ocr_list[i])
                        # 在 ocr_list 中删除该部分
                        ocr_list[i] = ""
                combined_text = ''.join(ocr_list)
                sentences = re.split(r'(?<=[。！？；;])', combined_text)
                sentences = [sentence.strip() for sentence in sentences if sentence.strip()]
                self.ocr_para.extend(sentences)
                self.ocr_para = pattern_para + self.ocr_para

                text_content = "\n".join(self.ocr_para)
                self.text_display.append(text_content)
                # 将当前状态保存到历史堆栈中
                self.update_history(self.text_display.toPlainText())
            
            return ImagePath, data

        except TypeError:
            if is_table_cell:
                print("未识别到表格，请尝试其他图像。")
            else:
                print("未识别到文字，请尝试其他图像。")
            return None, None
        except Exception as e:
            print(f"处理图像时出错: {e}")
            return None, None

    def process_ocr_text(self, ocr_list):
        if len(ocr_list) == 1:
            return ocr_list[0]
        
        processed_text = []
        for text in ocr_list:
            # 判断尾部的字符
            last_char = text[-1] if text else ''
            if text.endswith("\n"):
                prev_char = text[-2] if len(text) > 1 else ''
                if prev_char.isalpha():  # 英文字符
                    text = text[:-1] + " "  # 替换 "\n" 为 " "
                elif prev_char != "。":  # 不是中文句号
                    text = text[:-1]  # 移除 "\n"
            processed_text.append(text)
        
        return "".join(processed_text)
    
    def compare_with_reference(self, ocr_list):
        matched_results = []
        for ocr_text in ocr_list:
            if not should_match(ocr_text):
                continue
            if self.is_table_mode:
                matches = process.extract(ocr_text, self.table_list, limit=5, scorer=fuzz.token_sort_ratio)
            else:
                matches = process.extract(ocr_text, self.para_list, limit=5)
            matches.sort(key=lambda x: (abs(len(x[0]) - len(ocr_text)), -x[1]))
            best_match = matches[0]
            if best_match[1] < 30:
                best_match = (ocr_text, "不好的匹配", best_match[1])
            matched_results.append((ocr_text, best_match[0], best_match[1]))
        return matched_results

    def on_screenshot_done(self, exitCode, exitStatus):
        # 截图进程完成后，显示主界面
        self.show_main_window()

def is_already_running():
    program_name = "GaloisVerificationMutex"
    
    try:
        mutex = win32api.CreateMutex(None, 1, program_name)
        if ctypes.GetLastError() == win32api.ERROR_ALREADY_EXISTS:
            return True
    except:
        pass
    
    return False

if is_already_running():
    print("Another instance is already running.")
    sys.exit()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setStyleSheet(qdarkstyle.load_stylesheet_pyqt5())
    window = MainApp()
    window.show()
    # 进入程序事件循环
    sys.exit(app.exec_())