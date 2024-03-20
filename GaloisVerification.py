import sys
import keyboard
import qdarkstyle
from PyQt5.QtWidgets import (QApplication, QMainWindow, QSystemTrayIcon, QMenu, QAction, QLabel,
                            QPushButton, QTextBrowser, QCheckBox, QWidgetAction, QTextEdit, 
                            QVBoxLayout, QHBoxLayout, QWidget, QDesktopWidget, QFileDialog,
                            QShortcut)
from PyQt5.QtGui import QIcon, QPixmap, QKeySequence, QFont, QPalette, QColor
from PyQt5.QtCore import Qt, QProcess, QMetaObject, pyqtSlot, pyqtSignal
from PyQt5.QtWebEngineWidgets import QWebEngineView
from PPOCR_api import GetOcrApi
from PPOCR_visualize import visualize
from PIL import Image
from docx import Document
from fuzzywuzzy import fuzz, process
import difflib
import cv2
# import tbpu
import ujson as json
import re
import os
import numpy as np

def detect_lines(image, thickness=1, line_color=255):
    # 获取图像的高度和宽度
    pixel_height_line = image.shape[0]
    pixel_width_line = image.shape[1]

    # 记录横线的上下边界
    horizontal_lines = []

    # pixel_values = list(range(31))

    # 遍历图像的高度
    for height_x in range(0, pixel_height_line, thickness):
        count = 0
        # 检查横向连续的黑色像素长度
        for width_x in range(pixel_width_line - thickness + 1):
            # 如果当前位置及粗度范围内的像素都为黑色
            if np.all(np.any(image[height_x, width_x:width_x + thickness] == 0)):
                count += 1

            # 如果连续黑色像素长度超过当前粗度阈值的95%
            if count > int(pixel_width_line * 0.9):
                # 将该横向区域的像素值全部设为白色（255），去除横线
                image[height_x, :] = line_color
                # 记录横线的上下边界
                horizontal_lines.append([height_x, height_x + thickness - 1])
                break

    return image, horizontal_lines

def should_match(text):
    # 确保文本不包含英文逗号、英文句号或中文句号
    if not re.search(r'[.。]', text):
        return True
    return False

def find_similar_sentences(text1, text2, threshold=0.8):
    matches = []
    for sent1 in text1:
        for sent2 in text2:
            s = difflib.SequenceMatcher(None, sent1, sent2)
            if s.ratio() > threshold:
                matches.append((sent1, sent2))
    return matches

def extract_table_from_docx(docx_path):
    doc = Document(docx_path)
    extracted_content = []
    
    # 遍历文档中的表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extracted_content.append(cell.text.strip())
    
    return "\n".join(extracted_content)

def highlight_differences(text1, text2):
    d = difflib.ndiff(text1, text2)
    diff_list = list(d)
    
    result1 = []
    result2 = []
    for item in diff_list:
        if item[0] == " ":
            result1.append(item[2:])
            result2.append(item[2:])
        elif item[0] == "-":
            result1.append(f'<span style="background-color: #fd7bd4;">{item[2:]}</span>')
        elif item[0] == "+":
            result2.append(f'<span style="background-color: #55a6f7;">{item[2:]}</span>')
    
    return "".join(result1), "".join(result2)

scrollbar_style = """
<style>
    ::-webkit-scrollbar {
        width: 10px;  /* 设置滚动条宽度 */
    }
    ::-webkit-scrollbar-track {
        background: #19232d;  /* 设置滚动条轨道的颜色 */
        /* border: 1px solid #455364; */  /* 设置滚动条轨道的描边 */
        border-radius: 5px;   /* 设置滚动条轨道的圆角 */
        /* box-shadow: inset 0 0 5px grey; */
    }
    ::-webkit-scrollbar-thumb {
        background: #60798b;  /* 设置滚动条的颜色 */
        border-radius: 5px;   /* 设置滚动条的圆角 */
    }
    ::-webkit-scrollbar-thumb:hover {
        background: #346792;  /* 设置鼠标悬停时滚动条的颜色 */
    }
</style>
"""

class MainApp(QMainWindow):

    triggerScreenshot = pyqtSignal()  # 定义一个信号

    def __init__(self):
        super(MainApp, self).__init__()

        self.init_ui()
        self.init_tray_icon()
        self.init_hotkeys()

        self.triggerScreenshot.connect(self.take_screenshot)  # 连接信号到take_screenshot函数

        # 初始化历史堆栈和当前索引
        self.history = [""]
        self.current_index = 0

        # 表单模式
        self.is_table_mode = 0

    def init_ui(self):
        """初始化主要的用户界面组件"""
        self.setWindowTitle("Galois's Content Verification")
        self.setGeometry(100, 100, 1000, 800)
        # self.setFixedSize(1000, 800)
        self.center()

        # 创建小部件
        label_top = QLabel("———————————————————— Welcome to Galois's Content Verification machine for Product ————————————————————", self)
        font = QFont()
        font.setPointSize(11)
        font.setBold(True)
        font.setFamily("Arial")
        label_top.setFont(font)
        palette = label_top.palette()
        palette.setColor(QPalette.WindowText, QColor("#FF5733"))
        label_top.setPalette(palette)
        label_top.setAlignment(Qt.AlignCenter)

        self.results_browser = QWebEngineView(self)
        
        display_browser_font = QFont()
        display_browser_font.setPointSize(8)
        display_browser_font.setBold(True)
        display_browser_font.setFamily("Arial")

        vbox_display = QVBoxLayout()
        label_display = QLabel("OCR Display Area", self)
        label_display.setFont(display_browser_font)
        label_display.setStyleSheet("background-color: #101c28; border: 1px solid #555555; border-bottom: none; padding: 2px; border-top-left-radius: 10px; border-top-right-radius: 10px;")
        self.text_display = QTextBrowser(self)
        self.text_display.setStyleSheet("border-top: none; border-top-left-radius: 0px; border-top-right-radius: 0px;")
        vbox_display.addWidget(label_display)
        vbox_display.addWidget(self.text_display)
        self.button_comparison = QPushButton("文本校验", self)
        self.button_form_comparison = QPushButton("表单校验", self)

        hbox_buttons = QHBoxLayout()
        hbox_buttons.addWidget(self.button_comparison)
        hbox_buttons.addWidget(self.button_form_comparison)

        hbox_buttons.setSpacing(5)
        hbox_buttons.setContentsMargins(0, 5, 0, 5)

        vbox_display.addLayout(hbox_buttons)

        vbox_display.setSpacing(0)
        vbox_display.setContentsMargins(0, 0, 0, 0)
        
        button_select = QPushButton("选择参考文档", self)
        # 新建切换按钮并初始化为禁用表格模式
        self.toggle_table_mode_button = QPushButton("文本流模式", self)
        self.toggle_table_mode_button.setCheckable(True)

        hbox_select = QHBoxLayout()
        hbox_select.addWidget(self.toggle_table_mode_button)

        hbox_select.setSpacing(5)
        hbox_select.setContentsMargins(0, 5, 0, 5)

        hbox_select.addWidget(button_select)

        self.text_browser = QTextBrowser(self)
        self.text_browser.setStyleSheet("border-top: none; border-top-left-radius: 0px; border-top-right-radius: 0px;")
        # self.text_browser = QWebEngineView(self)
        
        vbox_browser = QVBoxLayout()
        label_browser = QLabel("Original Content", self)
        label_browser.setFont(display_browser_font)
        label_browser.setStyleSheet("background-color: #101c28; border: 1px solid #555555; border-bottom: none; padding: 2px; border-top-left-radius: 10px; border-top-right-radius: 10px;")
        vbox_browser.addWidget(label_browser)
        vbox_browser.addWidget(self.text_browser)
        vbox_browser.addLayout(hbox_select)

        vbox_browser.setSpacing(0)
        vbox_browser.setContentsMargins(0, 0, 0, 0)

        self.results_browser.page().setBackgroundColor(Qt.transparent)
        self.results_browser.setStyleSheet("border: none;")
        self.results_browser.setMinimumHeight(100)
        self.results_browser.setStyleSheet("background-color: #19232d;")

        # 布局
        hbox_inner = QHBoxLayout()
        hbox_inner.addLayout(vbox_display)
        hbox_inner.addLayout(vbox_browser)

        vbox_outer = QVBoxLayout()
        vbox_outer.addWidget(label_top)
        vbox_outer.addLayout(hbox_inner)
        vbox_outer.addWidget(self.results_browser, 1)

        central_widget = QWidget(self)
        central_widget.setLayout(vbox_outer)
        self.setCentralWidget(central_widget)

        self.setWindowIcon(QIcon("icon.ico"))
        label_top.setFocus()

        button_select.clicked.connect(self.select_reference_file)
        self.button_comparison.clicked.connect(self.compare_texts)
        self.button_form_comparison.clicked.connect(self.on_form_comparison_clicked)
        self.toggle_table_mode_button.toggled.connect(self.toggle_table_mode)

        # 默认状态
        self.button_comparison.setEnabled(False)
        self.button_form_comparison.setEnabled(False)
        self.toggle_table_mode_button.setEnabled(False)
        button_select.setEnabled(True)

    def init_tray_icon(self):
        """初始化系统托盘图标及其关联的操作和菜单。"""
        self.tray_icon = QSystemTrayIcon(self)
        self.tray_icon.setIcon(QIcon("icon.ico"))

        self.tray_menu = QMenu()
        self.show_action = QAction("显示主界面", self)
        self.hide_when_screenshot = QAction("截图时隐藏窗口", self)
        self.diff_show = QAction("仅显示差异内容", self)
        self.hide_when_screenshot.setCheckable(True)
        self.diff_show.setCheckable(True)
        self.diff_show.setChecked(True)
        self.quit_action = QAction("退出", self)

        self.show_action.triggered.connect(self.show_main_window)
        self.quit_action.triggered.connect(self.quit_app)

        self.tray_menu.addAction(self.show_action)
        self.tray_menu.addAction(self.hide_when_screenshot)
        self.tray_menu.addAction(self.diff_show)
        self.tray_menu.addSeparator()
        self.tray_menu.addAction(self.quit_action)

        self.tray_icon.setContextMenu(self.tray_menu)
        self.tray_icon.show()
        self.tray_icon.activated.connect(self.tray_icon_activated)

    def init_hotkeys(self):
        """使用keyboard模块初始化热键。"""
        keyboard.add_hotkey('win+shift+a', lambda: self.triggerScreenshot.emit())
        # keyboard.add_hotkey('win+shift+a', lambda: QMetaObject.invokeMethod(self, "take_screenshot", Qt.QueuedConnection))
        # keyboard.add_hotkey('ctrl+r', self.clear_text_display)
        # keyboard.add_hotkey('ctrl+r', lambda: QMetaObject.invokeMethod(self, "clear_text_display", Qt.QueuedConnection))
        self.shortcut_clear = QShortcut(QKeySequence("Ctrl+R"), self)
        self.shortcut_clear.activated.connect(self.clear_text_display)

        # 设置快捷键
        self.shortcut_undo = QShortcut(QKeySequence("Ctrl+Z"), self)
        self.shortcut_undo.activated.connect(self.undo)
        
        self.shortcut_redo = QShortcut(QKeySequence("Ctrl+Shift+Z"), self)
        self.shortcut_redo.activated.connect(self.redo)

        self.shortcut_clear_results = QShortcut(QKeySequence("Ctrl+L"), self)  # Ctrl+L 用于清除结果
        self.shortcut_clear_results.activated.connect(self.clear_results_browser)
    
    def toggle_table_mode(self, checked):
        if checked:
            self.toggle_table_mode_button.setText("表单模式")
            self.toggle_table_mode_button.setStyleSheet("background-color: #0066aa;")
            self.is_table_mode = 1
            self.button_comparison.setEnabled(False)
            self.button_form_comparison.setEnabled(True)
        else:
            self.toggle_table_mode_button.setText("文本流模式")
            self.toggle_table_mode_button.setStyleSheet("")
            self.is_table_mode = 0
            self.button_comparison.setEnabled(True)
            self.button_form_comparison.setEnabled(False)


    def clear_results_browser(self):
        # self.results_browser.clear()
        self.results_browser.setHtml("")

    def tray_icon_activated(self, reason):
        # 托盘图标双击时，显示主界面
        if reason == QSystemTrayIcon.DoubleClick:
            self.show_main_window()
    
    def center(self):
        # 获取屏幕的矩形和窗口的矩形
        screen = QDesktopWidget().screenGeometry()
        window = self.geometry()
        self.move(int((screen.width() - window.width()) / 2), int((screen.height() - window.height()) / 2))

    def show_main_window(self):
        # 激活并显示主界面
        self.setWindowState(self.windowState() & ~Qt.WindowMinimized | Qt.WindowActive)
        self.show()
        self.activateWindow()
        self.raise_()

    def quit_app(self):
        """退出应用程序"""
        # 关闭主窗口
        self.close()

        # 如果截图进程存在，结束该进程
        if hasattr(self, 'screenshot_process'):
            self.screenshot_process.terminate()
            self.screenshot_process.waitForFinished(2000)  # 等待2秒以给进程结束的时间

        # 隐藏托盘图标并退出应用
        self.tray_icon.hide()
        QApplication.quit()

    def closeEvent(self, event):
        # 重写窗口的关闭事件，使其隐藏窗口而不是退出应用
        event.ignore()
        self.hide()

    def update_history(self, text):
        # print("Before undo:", self.current_index, self.history)
        # 当文本更改时，将其添加到历史堆栈
        self.current_index += 1
        self.history = self.history[:self.current_index]
        self.history.append(text)
        # print("After undo:", self.current_index, self.history)
    
    def undo(self):
        # print("Before undo:", self.current_index, self.history)
        if self.current_index > 0:
            self.current_index -= 1
            self.text_display.setPlainText(self.history[self.current_index])
        # print("After undo:", self.current_index, self.history)

    def redo(self):
        # print("Before undo:", self.current_index, self.history)
        if self.current_index < len(self.history) - 1:
            self.current_index += 1
            self.text_display.setPlainText(self.history[self.current_index])
        # print("After undo:", self.current_index, self.history)
    
    def compare_texts(self):
        text_display_content = self.text_display.toPlainText()
        text_paras_content = "\n".join(item for item in self.para_list)
        text_display_list = text_display_content.split('\n')
        text_paras_list = text_paras_content.split('\n') + self.para_stream_list

        # from pprint import pprint
        # pprint(text_display_list)
        # pprint(text_paras_list)

        matches = find_similar_sentences(text_display_list, text_paras_list)

        result_texts = []
        for match in matches:
            diff_text1, diff_text2 = highlight_differences(match[0], match[1])
            if self.diff_show.isChecked():
                # 检查是否真的存在差异
                s = difflib.SequenceMatcher(None, diff_text1, diff_text2)
                if s.ratio() < 1.0:  # 如果文本不完全相同（即存在差异）
                    result_text = (f'<div style="color: white; font-size: 13px;">'
                                f'<hr style="border: 0; height: 1px; background-color: #22aacc;">'
                                f'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px;margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span> {diff_text1}<br>'
                                f'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px;margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span> {diff_text2}<br>'
                                f'</div>')
                    result_texts.append(result_text)
            else:
                result_text = (f'<div style="color: white; font-size: 13px;">'
                                f'<hr style="border: 0; height: 2px; background-color: #22aacc;">'
                                f'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px; margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span> {diff_text1}<br>'
                                f'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px; margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span> {diff_text2}<br>'
                                f'</div>')
                result_texts.append(result_text)

        final_html = '<html><body>' + scrollbar_style + ''.join(result_texts) + '</body></html>'
        self.results_browser.setHtml(final_html)
    
    def on_form_comparison_clicked(self):
        text_display_content = self.text_display.toPlainText()
        # 判断内容中是否包含换行符
        if "\n" in text_display_content:
            text_display_list = text_display_content.split("\n")
            matches = find_similar_sentences(text_display_list, self.table_list)
        else:
            matches = find_similar_sentences([text_display_content], self.table_list)

        result_texts = []
        for match in matches:
            diff_text1, diff_text2 = highlight_differences(match[0], match[1])
            if self.diff_show.isChecked():
                # 检查是否真的存在差异
                s = difflib.SequenceMatcher(None, diff_text1, diff_text2)
                if s.ratio() < 1.0:  # 如果文本不完全相同（即存在差异）
                    result_text = (f'<div style="color: white; font-size: 13px;">'
                                f'<hr style="border: 0; height: 2px; background-color: #22aacc;">'
                                f'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px; margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span> {diff_text1}<br>'
                                f'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px; margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span> {diff_text2}<br>'
                                f'</div>')
                    result_texts.append(result_text)
            else:
                result_text = (f'<div style="color: white; font-size: 13px;">'
                                f'<hr style="border: 0; height: 2px; background-color: #22aacc;">'
                                f'<span style="font-size: 11px; background-color: #fd7bd4; color: black; border-radius: 20px; margin-bottom: 2px; display: inline-block; font-weight: bold;">Design:</span> {diff_text1}<br>'
                                f'<span style="font-size: 11px; background-color: #55a6f7; color: black; border-radius: 20px; margin-top: 2px; display: inline-block; font-weight: bold;">Original:</span> {diff_text2}<br>'
                                f'</div>')
                result_texts.append(result_text)

        final_html = '<html><body>' + scrollbar_style + ''.join(result_texts) + '</body></html>'
        self.results_browser.setHtml(final_html)


    @pyqtSlot()
    def select_reference_file(self):
        try:
            options = QFileDialog.Options()
            fileName, _ = QFileDialog.getOpenFileName(self, "选择参考文档", "", "word文件 (*.docx);;所有文件 (*);;文本文件 (*.txt);;PDF 文件 (*.pdf)", options=options)
            
            if fileName:
                content_to_display = []
                self.para_list = []  # 保存符合正则表达式的段落内容
                self.table_list = []  # 保存表格的单元格内容
                self.para_stream_list = [] # 保存文本流内容

                # 如果是 .docx 文件
                if fileName.endswith(".docx"):
                    doc = Document(fileName)

                    pattern = r'^【.+?】(?:.+)'

                    para_exist = False
                    table_exist = False

                    for element in doc._element.body:
                        # 如果是段落
                        if element.tag.endswith('p'):
                            para = [p for p in doc.paragraphs if p._element is element][0]
                            if re.match(pattern, para.text):
                                self.para_list.append(para.text)
                            else:
                                self.para_stream_list.append(para.text)
                                # print(para.text)
                            if para.text.strip() != '':
                                content_to_display.append(para.text)
                                para_exist = True

                        # 如果是表格
                        elif element.tag.endswith('tbl'):
                            table = [t for t in doc.tables if t._element is element][0]
                            for row in table.rows:
                                for cell in row.cells:
                                    if cell.text.strip() != '':
                                        content_to_display.append(cell.text)
                                        self.table_list.append(cell.text)  # 保存表格单元格内容
                                        table_exist = True

                    if not para_exist:
                        print(f"{fileName}中不存在段落内容")
                    if not table_exist:
                        print(f"{fileName}中不存在表格")

                # 如果是其他类型的文件，例如 .txt
                else:
                    with open(fileName, 'r', encoding="utf-8") as file:
                        content_to_display.append(file.read())

                # 更新 self.text_browser
                self.text_browser.setText('\n'.join(content_to_display))

                if self.text_browser.toPlainText():
                    self.toggle_table_mode_button.setEnabled(True)
                    self.button_comparison.setEnabled(True)

        except Exception as e:
            print("无效的文件类型:", str(e))

    def clear_text_display(self):
        # 将当前状态保存到历史堆栈中
        self.update_history(self.text_display.toPlainText())
        self.text_display.clear()

    def take_screenshot(self):
        # 隐藏窗口，然后启动截图进程
        if self.hide_when_screenshot.isChecked():
            self.hide()
        self.start_screenshot_process()

    def start_screenshot_process(self):
        """启动截图进程的函数。"""

        # 创建一个QProcess对象，用于执行外部进程
        self.screenshot_process = QProcess(self)
        # 当有数据写入到标准输出时，连接到on_screenshot_output函数
        self.screenshot_process.readyReadStandardOutput.connect(self.on_screenshot_output)
        # 当截图进程完成时，连接到on_screenshot_done函数进行处理
        self.screenshot_process.finished.connect(self.on_screenshot_done)
        # 确定应用路径
        app_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))

        if getattr(sys, 'frozen', False):
            # 如果应用是打包的
            # screenshot_tool_path = os.path.join(app_path, 'screenshot_tool')
            screenshot_tool_path = os.path.join(app_path, 'screenshot_tool')
            screenshot_tool_path = f'"{screenshot_tool_path}"'  # 将路径用引号括起来
            self.screenshot_process.start(screenshot_tool_path)
        else:
            # 在开发环境中
            self.screenshot_process.start('python', [os.path.join(app_path, 'screenshot_tool.py')])

    def on_screenshot_output(self):
        """当screenshot_tool.py有标准输出时的处理函数"""

        # 从进程的标准输出中读取文件名
        filename = bytes(self.screenshot_process.readAllStandardOutput()).decode().strip()
        # print("Screenshot saved to:", filename)
        filename = f"{os.path.dirname(os.path.abspath(__file__))}\\{filename}"
        
        # 根据截图内容类型调用ppocr_to_data
        if self.is_table_mode:  # 判断当前是在处理表格模式还是文本模式
            self.ppocr_to_data(filename=filename, is_table_cell=True)
        else:
            self.ppocr_to_data(filename=filename, is_table_cell=False)
    
    def ppocr_to_data(self, filename, is_table_cell=False):
        try:
            ImagePath = filename
            # 初始化识别器
            ocr = GetOcrApi(os.path.abspath(os.path.join("ppocr", "PaddleOCR-json.exe"))) # 开发环境使用

            # app_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
            # ocr_exe_path = os.path.join(app_path, "ppocr", "PaddleOCR-json.exe")
            # ocr = GetOcrApi(ocr_exe_path) # 打包环境使用

            # print(f'初始化OCR成功，进程号为{ocr.ret.pid}')
            # print(f'\n图片路径：{ImagePath}')
            res = ocr.run(ImagePath)
            data = res['data']

            ocr_list = [item['text'] for item in data]
            
            if is_table_cell:  # 如果是表格单元格的内容
                processed_text = self.process_ocr_text(ocr_list)
                # 与self.table_list中的内容进行比对
                matches = self.compare_with_reference([processed_text])
                # TODO: 根据matches进行进一步的处理，比如高亮等
                self.text_display.append(processed_text)  # 假设显示表格内容的控件叫做text_display
                # from pprint import pprint
                # pprint(matches)
            
            else:
                text_browser_content = self.text_browser.toPlainText()
                text_browser_list = text_browser_content.split('\n')

                matched_results = self.compare_with_reference(ocr_list)
                matched_texts = [item[0] for item in matched_results]
                filtered_list = [text for text in ocr_list if text not in matched_texts]
                merged_list = []
                current_chunk = ""
                pattern = r'^【.+?】(?:.+)'
                for text in filtered_list:
                    if current_chunk and current_chunk[-1] == '。':
                        # 如果当前块以句号结尾，则将其添加到merged_list中并清空current_chunk
                        merged_list.append(current_chunk)
                        current_chunk = text
                    elif re.match(pattern, current_chunk):
                        merged_list.append(current_chunk)
                        current_chunk = text
                    else:
                        # 否则，将text追加到current_chunk中
                        current_chunk += text

                # 添加最后的块（如果存在）
                if current_chunk:
                    merged_list.append(current_chunk)

                text_content = "\n".join(item for item in matched_texts + merged_list)
                self.text_display.append(text_content)
                # 将当前状态保存到历史堆栈中
                self.update_history(self.text_display.toPlainText())

        except TypeError:
            print("未识别到文字，请尝试其他图像。")
        except Exception as e:
            print(f"处理图像时出错: {e}")

    def process_ocr_text(self, ocr_list):
        if len(ocr_list) == 1:
            return ocr_list[0]
        
        processed_text = []
        for text in ocr_list:
            # 判断尾部的字符
            last_char = text[-1] if text else ''
            if text.endswith("\n"):
                prev_char = text[-2] if len(text) > 1 else ''
                if prev_char.isalpha():  # 英文字符
                    text = text[:-1] + " "  # 替换 "\n" 为 " "
                elif prev_char != "。":  # 不是中文句号
                    text = text[:-1]  # 移除 "\n"
            processed_text.append(text)
        
        return "".join(processed_text)
    
    def compare_with_reference(self, ocr_list):
        matched_results = []
        for ocr_text in ocr_list:
            if not should_match(ocr_text):
                continue
            if self.is_table_mode:
                matches = process.extract(ocr_text, self.table_list, limit=5)
            else:
                matches = process.extract(ocr_text, self.para_list, limit=5)
            matches.sort(key=lambda x: (abs(len(x[0]) - len(ocr_text)), -x[1]))
            best_match = matches[0]
            if best_match[1] < 30:
                best_match = (ocr_text, "不好的匹配", best_match[1])
            matched_results.append((ocr_text, best_match[0], best_match[1]))
        return matched_results

    def on_screenshot_done(self, exitCode, exitStatus):
        # 截图进程完成后，显示主界面
        self.show_main_window()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setStyleSheet(qdarkstyle.load_stylesheet_pyqt5())
    window = MainApp()
    window.show()
    # 进入程序事件循环
    sys.exit(app.exec_())