import sys
import keyboard
import qdarkstyle
from PyQt5.QtWidgets import (QApplication, QMainWindow, QSystemTrayIcon, QMenu, QAction, QLabel,
                            QPushButton, QTextBrowser, QCheckBox, QWidgetAction, QTextEdit, 
                            QVBoxLayout, QHBoxLayout, QWidget, QDesktopWidget, QFileDialog,
                            QShortcut)
from PyQt5.QtGui import QIcon, QPixmap, QKeySequence, QFont
from PyQt5.QtCore import Qt, QProcess, QMetaObject, pyqtSlot, pyqtSignal
from PyQt5.QtWebEngineWidgets import QWebEngineView
from PPOCR_api import GetOcrApi
from PPOCR_visualize import visualize
from PIL import Image
from docx import Document
import difflib
import cv2
import tbpu
import ujson as json
import re
import os


def merge_text_blocks(textBlocks, image_width):
    sorted_text_blocks = sorted(textBlocks, key=lambda x: x['box'][0][1])
    merged_texts = []
    current_line = ""
    for i, block in enumerate(sorted_text_blocks):
        if i > 0 and (block['box'][0][1] - sorted_text_blocks[i-1]['box'][3][1]) > 5:
            merged_texts.append(current_line.strip())
            current_line = block['text']
        else:
            current_line += " " + block['text']
        if (block['box'][2][0] < image_width * 0.9) or block['text'].endswith(('。', '！', '？')):
            merged_texts.append(current_line.strip())
            current_line = ""
    if current_line:
        merged_texts.append(current_line.strip())
    
    # 为合并后的段落计算边界框
    merged_boxes = []
    start_idx = 0
    for i, block in enumerate(sorted_text_blocks):
        if i > 0 and (block['box'][0][1] - sorted_text_blocks[i-1]['box'][3][1]) > 5:
            box = get_bounding_box(sorted_text_blocks[start_idx:i])
            merged_boxes.append(box)
            start_idx = i

    if start_idx < len(sorted_text_blocks):
        box = get_bounding_box(sorted_text_blocks[start_idx:])
        merged_boxes.append(box)
    
    return "\n".join(merged_texts), merged_boxes

def get_bounding_box(blocks):
    min_x = min([block['box'][0][0] for block in blocks])
    min_y = min([block['box'][0][1] for block in blocks])
    max_x = max([block['box'][2][0] for block in blocks])
    max_y = max([block['box'][2][1] for block in blocks])
    return ((min_x, min_y), (max_x, max_y))

def find_similar_sentences(text1, text2, threshold=0.3):
    # 分割文本为句子
    sentences1 = split_into_sentences(text1)
    sentences2 = split_into_sentences(text2)
    # from pprint import pprint
    # pprint(sentences1)
    # print(sentences2)
    # exit()

    matches = []

    for sent1 in sentences1:
        for sent2 in sentences2:
            s = difflib.SequenceMatcher(None, sent1, sent2)
            if s.ratio() > threshold:
                matches.append((sent1, sent2))
    return matches

def split_into_sentences(text):
    """
    Split the text into sentences based on punctuation.
    """
    # sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?|\!)\s', text)
    sentences = re.split(r'(?<=[。？！])\s*', text)
    return [sent.replace('\n', '').strip() for sent in sentences if sent]

def extract_table_from_docx(docx_path):
    doc = Document(docx_path)
    extracted_content = []
    
    # 遍历文档中的表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extracted_content.append(cell.text.strip())
    
    return "\n".join(extracted_content)

def highlight_differences(text1, text2):
    d = difflib.ndiff(text1, text2)
    diff_list = list(d)
    
    result1 = []
    result2 = []
    for item in diff_list:
        if item[0] == " ":
            result1.append(item[2:])
            result2.append(item[2:])
        elif item[0] == "-":
            result1.append(f'<span style="background-color: red;">{item[2:]}</span>')
        elif item[0] == "+":
            result2.append(f'<span style="background-color: green;">{item[2:]}</span>')
    
    return "".join(result1), "".join(result2)

class MainApp(QMainWindow):

    triggerScreenshot = pyqtSignal()  # 定义一个信号

    def __init__(self):
        super(MainApp, self).__init__()

        self.init_ui()
        self.init_tray_icon()
        self.init_hotkeys()

        self.triggerScreenshot.connect(self.take_screenshot)  # 连接信号到take_screenshot函数

        # 初始化历史堆栈和当前索引
        self.history = [""]
        self.current_index = 0

    def init_ui(self):
        """初始化主要的用户界面组件。"""
        self.setWindowTitle("Galois's Description Detection")
        self.setGeometry(100, 100, 1000, 800)
        # self.setFixedSize(1000, 800)
        self.center()

        # 创建小部件
        label_top = QLabel("Welcome to Galois's description detection tool for Product", self)
        font = QFont()
        font.setBold(True)
        label_top.setFont(font)
        label_top.setAlignment(Qt.AlignCenter)
        # self.results_browser = QTextBrowser(self)
        self.results_browser = QWebEngineView(self)
        
        # 新的垂直布局，包含一个QTextBrowser和一个按钮
        vbox_left = QVBoxLayout()
        self.text_display = QTextBrowser(self)
        vbox_left.addWidget(self.text_display)
        button_comparison = QPushButton("文本比对", self)
        vbox_left.addWidget(button_comparison)
        
        button_select = QPushButton("选择参考文档", self)
        self.text_browser = QTextBrowser(self)
        # self.text_browser = QWebEngineView(self)
        
        vbox_right = QVBoxLayout()
        vbox_right.addWidget(button_select)
        vbox_right.addWidget(self.text_browser)
        self.results_browser.page().setBackgroundColor(Qt.transparent)
        self.results_browser.setStyleSheet("border: none;")
        self.results_browser.setMinimumHeight(100)
        self.results_browser.setStyleSheet("background-color: #19232d;")

        
        # 布局
        hbox_inner = QHBoxLayout()
        hbox_inner.addLayout(vbox_left)
        hbox_inner.addLayout(vbox_right)

        vbox_outer = QVBoxLayout()
        vbox_outer.addWidget(label_top)
        vbox_outer.addLayout(hbox_inner)
        vbox_outer.addWidget(self.results_browser, 1)

        central_widget = QWidget(self)
        central_widget.setLayout(vbox_outer)
        self.setCentralWidget(central_widget)

        self.setWindowIcon(QIcon("icon.png"))
        label_top.setFocus()

        button_select.clicked.connect(self.select_reference_file)
        button_comparison.clicked.connect(self.compare_texts)


    def init_tray_icon(self):
        """初始化系统托盘图标及其关联的操作和菜单。"""
        self.tray_icon = QSystemTrayIcon(self)
        self.tray_icon.setIcon(QIcon("icon.png"))

        self.tray_menu = QMenu()
        self.show_action = QAction("显示主界面", self)
        self.hide_when_screenshot = QAction("截图时隐藏窗口", self)
        self.hide_when_screenshot.setCheckable(True)
        self.quit_action = QAction("退出", self)

        self.show_action.triggered.connect(self.show_main_window)
        self.quit_action.triggered.connect(self.quit_app)

        self.tray_menu.addAction(self.show_action)
        self.tray_menu.addAction(self.hide_when_screenshot)
        self.tray_menu.addSeparator()
        self.tray_menu.addAction(self.quit_action)

        self.tray_icon.setContextMenu(self.tray_menu)
        self.tray_icon.show()
        self.tray_icon.activated.connect(self.tray_icon_activated)

    def init_hotkeys(self):
        """使用keyboard模块初始化热键。"""
        keyboard.add_hotkey('win+shift+a', lambda: self.triggerScreenshot.emit())
        # keyboard.add_hotkey('win+shift+a', lambda: QMetaObject.invokeMethod(self, "take_screenshot", Qt.QueuedConnection))
        # keyboard.add_hotkey('ctrl+r', self.clear_text_display)
        # keyboard.add_hotkey('ctrl+r', lambda: QMetaObject.invokeMethod(self, "clear_text_display", Qt.QueuedConnection))
        self.shortcut_clear = QShortcut(QKeySequence("Ctrl+R"), self)
        self.shortcut_clear.activated.connect(self.clear_text_display)

        # 设置快捷键
        self.shortcut_undo = QShortcut(QKeySequence("Ctrl+Z"), self)
        self.shortcut_undo.activated.connect(self.undo)
        
        self.shortcut_redo = QShortcut(QKeySequence("Ctrl+Shift+Z"), self)
        self.shortcut_redo.activated.connect(self.redo)

        self.shortcut_clear_results = QShortcut(QKeySequence("Ctrl+L"), self)  # Ctrl+L 用于清除结果
        self.shortcut_clear_results.activated.connect(self.clear_results_browser)

    def clear_results_browser(self):
        # self.results_browser.clear()
        self.results_browser.setHtml("")

    def tray_icon_activated(self, reason):
        # 托盘图标双击时，显示主界面
        if reason == QSystemTrayIcon.DoubleClick:
            self.show_main_window()
    
    def center(self):
        # 获取屏幕的矩形和窗口的矩形
        screen = QDesktopWidget().screenGeometry()
        window = self.geometry()
        self.move(int((screen.width() - window.width()) / 2), int((screen.height() - window.height()) / 2))

    def show_main_window(self):
        # 激活并显示主界面
        self.setWindowState(self.windowState() & ~Qt.WindowMinimized | Qt.WindowActive)
        self.show()
        self.activateWindow()
        self.raise_()

    def quit_app(self):
        """退出应用程序"""
        # 关闭主窗口
        self.close()

        # 如果截图进程存在，结束该进程
        if hasattr(self, 'screenshot_process'):
            self.screenshot_process.terminate()
            self.screenshot_process.waitForFinished(2000)  # 等待2秒以给进程结束的时间

        # 隐藏托盘图标并退出应用
        self.tray_icon.hide()
        QApplication.quit()

    def closeEvent(self, event):
        # 重写窗口的关闭事件，使其隐藏窗口而不是退出应用
        event.ignore()
        self.hide()

    def update_history(self, text):
        # print("Before undo:", self.current_index, self.history)
        # 当文本更改时，将其添加到历史堆栈
        self.current_index += 1
        self.history = self.history[:self.current_index]
        self.history.append(text)
        # print("After undo:", self.current_index, self.history)
    
    def undo(self):
        # print("Before undo:", self.current_index, self.history)
        if self.current_index > 0:
            self.current_index -= 1
            self.text_display.setPlainText(self.history[self.current_index])
        # print("After undo:", self.current_index, self.history)

    def redo(self):
        # print("Before undo:", self.current_index, self.history)
        if self.current_index < len(self.history) - 1:
            self.current_index += 1
            self.text_display.setPlainText(self.history[self.current_index])
        # print("After undo:", self.current_index, self.history)
    
    def compare_texts(self):
        text_display_content = self.text_display.toPlainText()
        text_browser_content = self.text_browser.toPlainText()

        matches = find_similar_sentences(text_display_content, text_browser_content)

        result_texts = []
        for match in matches:
            # print(f"Match found:\nFrom text_display: {match[0]}\nFrom text_browser: {match[1]}\n")
            # result_text = f"匹配项:\n设计图: {match[0]}\n参照原本: {match[1]}\n"
            # result_text = (f'<table width="100%" cellspacing="0" cellpadding="0" border="0">'
            #            f'<tr><td style="background-color: #e6e6e6;">匹配项:</td></tr>'
            #            f'</table>'
            #            f'<span style="background-color: #99ccff; border-radius: 20px; padding: 2px;">设计图:</span> {match[0]}<br>'
            #            f'<span style="background-color: #99ccff; border-radius: 20px; padding: 2px;">参照原本:</span> {match[1]}<br>')
            # result_texts.append(result_text)
            diff_text1, diff_text2 = highlight_differences(match[0], match[1])
            result_text = (f'<div style="color: white; font-size: 14px;">'
                        f'<hr style="border: 0; height: 2px; background-color: #22cccc;">'
                        f'<span style="font-size: 12px; background-color: #ffcc66; color: black; border-radius: 20px; padding: 2px; margin-bottom: 2px; display: inline-block;">设计图:</span> {diff_text1}<br>'
                        f'<span style="font-size: 12px; background-color: #99ccff; color: black; border-radius: 20px; padding: 2px; margin-top: 2px; display: inline-block;">参照原本:</span> {diff_text2}<br>'
                        f'</div>')
            result_texts.append(result_text)

        final_html = '<html><body>' + ''.join(result_texts) + '</body></html>'
        self.results_browser.setHtml(final_html)

    @pyqtSlot()
    def select_reference_file(self):
        try:
            options = QFileDialog.Options()
            fileName, _ = QFileDialog.getOpenFileName(self, "选择参考文档", "", "word文件 (*.docx);;所有文件 (*);;文本文件 (*.txt);;PDF 文件 (*.pdf)", options=options)
            
            if fileName:
                content_to_display = []

                # 如果是 .docx 文件
                if fileName.endswith(".docx"):
                    doc = Document(fileName)

                    para_exist = False
                    table_exist = False

                    for element in doc._element.body:
                        # 如果是段落
                        if element.tag.endswith('p'):
                            para = [p for p in doc.paragraphs if p._element is element][0]
                            if para.text.strip() != '':
                                content_to_display.append(para.text)
                                para_exist = True

                        # 如果是表格
                        elif element.tag.endswith('tbl'):
                            table = [t for t in doc.tables if t._element is element][0]
                            for row in table.rows:
                                for cell in row.cells:
                                    if cell.text.strip() != '':
                                        content_to_display.append(cell.text)
                                        table_exist = True

                    if not para_exist:
                        print(f"{fileName}中不存在段落内容")
                    if not table_exist:
                        print(f"{fileName}中不存在表格")

                # 如果是其他类型的文件，例如 .txt
                else:
                    with open(fileName, 'r', encoding="utf-8") as file:
                        content_to_display.append(file.read())

                # 更新 self.text_browser
                self.text_browser.setText('\n'.join(content_to_display))
        except Exception as e:
            print("无效的文件类型:", str(e))


    def clear_text_display(self):
        # 将当前状态保存到历史堆栈中
        self.update_history(self.text_display.toPlainText())
        self.text_display.clear()

    def take_screenshot(self):
        # 隐藏窗口，然后启动截图进程
        if self.hide_when_screenshot.isChecked():
            self.hide()
        self.start_screenshot_process()

    def start_screenshot_process(self):
        """启动截图进程的函数。"""

        # 创建一个QProcess对象，用于执行外部进程
        self.screenshot_process = QProcess(self)
        # 当有数据写入到标准输出时，连接到on_screenshot_output函数
        self.screenshot_process.readyReadStandardOutput.connect(self.on_screenshot_output)
        # 当截图进程完成时，连接到on_screenshot_done函数进行处理
        self.screenshot_process.finished.connect(self.on_screenshot_done)
        # 使用QProcess启动外部的Python进程并执行screenshot_tool.py脚本
        self.screenshot_process.start('python', ['screenshot_tool.py'])
    
    def on_screenshot_output(self):
        """当screenshot_tool.py有标准输出时的处理函数"""

        # 从进程的标准输出中读取文件名
        filename = bytes(self.screenshot_process.readAllStandardOutput()).decode().strip()
        # print("Screenshot saved to:", filename)
        filename = f"{os.path.dirname(os.path.abspath(__file__))}\\{filename}"
        self.ppocr_to_data(filename=filename)
    
    def ppocr_to_data(self, filename):
        try:
            ImagePath = filename
            # 初始化识别器
            ocr = GetOcrApi(os.path.abspath(os.path.join("ppocr", "PaddleOCR-json.exe")))

            # print(f'初始化OCR成功，进程号为{ocr.ret.pid}')
            # print(f'\n图片路径：{ImagePath}')
            res = ocr.run(ImagePath)
            data = res['data']

            # # 获取图像宽度
            # with Image.open(filename) as img:
            #     image_width, _ = img.size
            
            # # 使用merge_text_blocks函数合并文本块
            # merged_text_content, merged_boxes = merge_text_blocks(data, image_width)

            # # 将合并后的文本显示在QTextEdit中
            # self.text_display.append(merged_text_content)

            # # 绘制框到图像上
            # img = cv2.imread(filename)
            # for box in merged_boxes:
            #     cv2.rectangle(img, box[0], box[1], (0, 255, 0), 2) # 绿色框，线宽2
            # # 显示图像
            # cv2.imshow('Merged Paragraphs', img)
            # cv2.waitKey(0)
            # cv2.destroyAllWindows()

            # with open(f"{filename[:-4]}_data.txt", "w") as file:
            #     json.dump(data, file)
            
            # visualize(data, ImagePath).show()
            img1 = visualize(data, ImagePath).get(isOrder=True)
            textBlocksNew = tbpu.run_merge_line_h_m_paragraph(data)
            # img2 = visualize(textBlocksNew, ImagePath).get(isOrder=True)
            # # visualize.createContrast(img1, img2).show()

            # # 将textBlocksNew['data']内容显示在QTextEdit中
            text_content = "\n".join(item['text'] for item in textBlocksNew)
            # text_content = "\n".join(item['text'] for item in data)
            # # self.text_edit.setText(text_content)
            self.text_display.append(text_content)

            # 将当前状态保存到历史堆栈中
            self.update_history(self.text_display.toPlainText())

            # with open(f"{filename[:-4]}_merge.txt", "w") as file:
            #     json.dump(text_content, file)
            # with open(f"{filename[:-4]}_merge.txt", "w", encoding="utf-8") as file:
            #     json.dump(textBlocksNew, file, ensure_ascii=False, indent=4)
        except TypeError:
            print("未识别到文字，请尝试其他图像。")
        except Exception as e:
            print(f"处理图像时出错: {e}")


    def on_screenshot_done(self, exitCode, exitStatus):
        # 截图进程完成后，显示主界面
        self.show_main_window()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setStyleSheet(qdarkstyle.load_stylesheet_pyqt5())
    window = MainApp()
    window.show()
    # 进入程序事件循环
    sys.exit(app.exec_())